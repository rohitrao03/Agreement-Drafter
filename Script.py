import tkinter as tk
from tkinter import ttk
from tkinter import*
from docx import Document
import inflect



class ScrollableFrame(ttk.Frame):
    def __init__(self, container, *args, **kwargs):
        super().__init__(container, *args, **kwargs)
        canvas = tk.Canvas(self,width=1250, height=800,bg='#f1f1f1')
        scrollbar = ttk.Scrollbar(self, orient="vertical", command=canvas.yview)
        self.scrollable_frame = ttk.Frame(canvas)

        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(
                scrollregion=canvas.bbox("all")
            )
        )

        canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw",width=1250, height=1500)

        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
    
        
        
root = tk.Tk()
root.configure(bg='#f1f1f1')


#adding window name and size
root.title("Generate Agreement")
root.geometry('1250x800')

frame = ScrollableFrame(root)



#adding label
lbl = Label(frame.scrollable_frame, text="Rent Agreement", font=("Arial Bold", 30),bg='#f1f1f1')
lbl.grid(column=2, row=0)

#line seperations    
Label(frame.scrollable_frame, text='                           ',font=('Arial',15),bg='#f1f1f1').grid(column=1,row=2)



#adding text box and getting text for city and state[PLACE]
Label(frame.scrollable_frame, text='City',font=('Arial',15),bg='#f1f1f1').grid(column=1,row=4)
city = Entry(frame.scrollable_frame,width=25)
city.grid(column=2, row=4)
Label(frame.scrollable_frame, text='State',font=('Arial',15),bg='#f1f1f1').grid(column=1,row=7)
state = Entry(frame.scrollable_frame,width=25)
state.grid(column=2,row=7)
#adding date[]

Label(frame.scrollable_frame, text='Enter Date(dd/mm/yyyy)',font=('Arial',15),bg='#f1f1f1').grid(column=1,row=10)
date = Entry(frame.scrollable_frame,width=25)
date.grid(column=2,row=10)

#adding the number of landlords
Label(frame.scrollable_frame, text='Number of Landlords',font=('Arial',15),bg='#f1f1f1').grid(column=1,row=13)
spin = Spinbox(frame.scrollable_frame, from_=0, to=100, width=5)
spin.grid(column=2,row=13)

def get_landlords_details():
    global numb_landlords
    numb_landlords=spin.get()
    lis_landlords_name=[0]*int(numb_landlords)
    lis_landlords_adhcin=[0]*int(numb_landlords)
    lis_landlords_add=[0]*int(numb_landlords)
    for i in range(int(numb_landlords)):
        Label(frame.scrollable_frame, text='Name of Landlord',font=('Arial',15),bg='#f1f1f1').grid(column=1,row=15+i)
        lis_landlords_name[i] = Entry(frame.scrollable_frame,width=25)
        lis_landlords_name[i].grid(column=2,row=15+i)
        Label(frame.scrollable_frame, text='Adhaar no./CIN',font=('Arial',15),bg='#f1f1f1').grid(column=3,row=15+i)
        lis_landlords_adhcin[i] = Entry(frame.scrollable_frame,width=25)
        lis_landlords_adhcin[i].grid(column=4,row=15+i)
        Label(frame.scrollable_frame, text='Address',font=('Arial',15),bg='#f1f1f1').grid(column=5,row=15+i)
        lis_landlords_add[i] = Entry(frame.scrollable_frame,width=25)
        lis_landlords_add[i].grid(column=6,row=15+i)
    def generate_landlords_details():
        lis_landlords_name_new = [0]*int(numb_landlords)
        lis_landlords_adhcin_new = [0]*int(numb_landlords)
        lis_landlords_add_new = [0]*int(numb_landlords)
        for i in range(int(numb_landlords)):
            lis_landlords_name_new[i] = lis_landlords_name[i].get()
            lis_landlords_adhcin_new[i] = lis_landlords_adhcin[i].get()
            lis_landlords_add_new[i] = lis_landlords_add[i].get()
        lis_name_adhcin_add=[]
        global landlord_details_final_global
        for i in range(int(numb_landlords)):
            lis_name_adhcin_add.append(('Name: '+lis_landlords_name_new[i]+'    Adhaar no./CIN: ' +lis_landlords_adhcin_new[i] + '     Address: '+ lis_landlords_add_new[i]))
        landlord_details_final_global='\n'.join(lis_name_adhcin_add)
    btn2 = Button(frame.scrollable_frame, text="Enter Tenants details", bg="black", fg="white",command=lambda:[generate_landlords_details(),get_tenants_details()])
    btn2.grid(column=2, row=26)
    
#adding the number of tenants
Label(frame.scrollable_frame, text='Number of Tenants',font=('Arial',15),bg='#f1f1f1').grid(column=1,row=25)
spin1 = Spinbox(frame.scrollable_frame, from_=0, to=100, width=5)
spin1.grid(column=2,row=25)
    
def get_tenants_details():
    global numb_tenants
    numb_tenants=spin1.get()
    lis_tenants_name=[0]*int(numb_tenants)
    lis_tenants_adhcin=[0]*int(numb_tenants)
    lis_tenants_add=[0]*int(numb_tenants)
    for i in range(int(numb_tenants)):
        Label(frame.scrollable_frame, text='Name of Tenant',font=('Arial',15),bg='#f1f1f1').grid(column=1,row=27+i)
        lis_tenants_name[i] = Entry(frame.scrollable_frame,width=25)
        lis_tenants_name[i].grid(column=2,row=27+i)
        Label(frame.scrollable_frame, text='Adhaar no./CIN',font=('Arial',15),bg='#f1f1f1').grid(column=3,row=27+i)
        lis_tenants_adhcin[i] = Entry(frame.scrollable_frame,width=25)
        lis_tenants_adhcin[i].grid(column=4,row=27+i)
        Label(frame.scrollable_frame, text='Address',font=('Arial',15),bg='#f1f1f1').grid(column=5,row=27+i)
        lis_tenants_add[i] = Entry(frame.scrollable_frame,width=25)
        lis_tenants_add[i].grid(column=6,row=27+i)
    def generate_tenants_details():
        lis_tenants_name_new = [0]*int(numb_tenants)
        lis_tenants_adhcin_new = [0]*int(numb_tenants)
        lis_tenants_add_new = [0]*int(numb_tenants)
        for i in range(int(numb_tenants)):
            lis_tenants_name_new[i] = lis_tenants_name[i].get()
            lis_tenants_adhcin_new[i] = lis_tenants_adhcin[i].get()
            lis_tenants_add_new[i] = lis_tenants_add[i].get()
        lis_name_adhcin_add=[]
        global tenant_details_final_global
        for i in range(int(numb_tenants)):
            lis_name_adhcin_add.append(('Name: '+lis_tenants_name_new[i]+'    Adhaar no./CIN: ' +lis_tenants_adhcin_new[i] + '     Address: '+ lis_tenants_add_new[i]))
        tenant_details_final_global='\n'.join(lis_name_adhcin_add)
    btn3 = Button(frame.scrollable_frame, text="Save Tenants details", bg="black", fg="white",command=lambda:generate_tenants_details())
    btn3.grid(column=3, row=37)


#line seperations    
Label(frame.scrollable_frame, text='                           ',font=('Arial',15),bg='#f1f1f1').grid(column=1,row=38)
Label(frame.scrollable_frame, text='                           ',font=('Arial',15),bg='#f1f1f1').grid(column=2,row=38)


#property address
Label(frame.scrollable_frame, text='Property address',font=('Arial',15),bg='#f1f1f1').grid(column=1,row=40)
prop = Entry(frame.scrollable_frame,width=40)
prop.grid(column=2,row=40)

#line seperations    
Label(frame.scrollable_frame, text='                           ',font=('Arial',15),bg='#f1f1f1').grid(column=1,row=41)

#rent amount and abel monthly/quarterly/yearly
Label(frame.scrollable_frame, text='Select rent period',font=('Arial',15),bg='#f1f1f1').grid(column=1,row=42)
OptionList=[" monthly "," quarterly ","   yearly   "]
var1 = StringVar(frame.scrollable_frame)
var1.set('Select')

Lb = OptionMenu(frame.scrollable_frame ,var1, *OptionList)
Lb.config(width=15, font=('Arial', 15))
Lb.grid(column=2,row=42)


#line seperations    
Label(frame.scrollable_frame, text='                           ',font=('Arial',15),bg='#f1f1f1').grid(column=1,row=43)

#making yearly/monthly/quarterly label
def callback(*args):
    global var2
    Label(frame.scrollable_frame,text='Enter'+var1.get()+'rent amount',font=('Arial', 15), bg='#f1f1f1').grid(column=1,row=44)
    var2=var1.get().strip()

var1.trace("w", callback)

rent_am = Entry(frame.scrollable_frame,width=30)
rent_am.grid(column=2,row=44)

#line seperations    
Label(frame.scrollable_frame, text='                           ',font=('Arial',15),bg='#f1f1f1').grid(column=1,row=43)


#duration of contract
Label(frame.scrollable_frame, text='Duration of contract(months)',font=('Arial',15),bg='#f1f1f1').grid(column=1,row=46)
duration = Entry(frame.scrollable_frame,width=5)
duration.grid(column=2,row=46)


#line seperations    
Label(frame.scrollable_frame, text='                           ',font=('Arial',15),bg='#f1f1f1').grid(column=1,row=47)

#purpose input and drop box
Label(frame.scrollable_frame, text='Choose purpose',font=('Arial',15),bg='#f1f1f1').grid(column=1,row=48)
purpose_list=["residential",'commercial']
var3 = StringVar(frame.scrollable_frame)
var3.set('choose purpose')


purp = OptionMenu(frame.scrollable_frame ,var3, *purpose_list)
purp.config(width=15, font=('Arial', 15))
purp.grid(column=2,row=48)
def caller(*args):
    global var4
    var4=var3.get()

var3.trace("w",caller)


#line seperations    
Label(frame.scrollable_frame, text='                           ',font=('Arial',15),bg='#f1f1f1').grid(column=1,row=49)

#expiry date of agreement


Label(frame.scrollable_frame, text='Enter expiry Date(dd/mm/yyyy)',font=('Arial',15),bg='#f1f1f1').grid(column=1,row=50)
exp_date = Entry(frame.scrollable_frame,width=25)
exp_date.grid(column=2,row=50)

Label(frame.scrollable_frame, text='                           ',font=('Arial',15),bg='#f1f1f1').grid(column=1,row=51)
Label(frame.scrollable_frame, text='                           ',font=('Arial',15),bg='#f1f1f1').grid(column=1,row=52)
Label(frame.scrollable_frame, text='                           ',font=('Arial',15),bg='#f1f1f1').grid(column=1,row=53)
Label(frame.scrollable_frame, text='                           ',font=('Arial',15),bg='#f1f1f1').grid(column=1,row=54)
Label(frame.scrollable_frame, text='                           ',font=('Arial',15),bg='#f1f1f1').grid(column=1,row=55)
Label(frame.scrollable_frame, text='                           ',font=('Arial',15),bg='#f1f1f1').grid(column=1,row=56)
Label(frame.scrollable_frame, text='                           ',font=('Arial',15),bg='#f1f1f1').grid(column=1,row=57)
Label(frame.scrollable_frame, text='                           ',font=('Arial',15),bg='#f1f1f1').grid(column=1,row=58)
Label(frame.scrollable_frame, text='                           ',font=('Arial',15),bg='#f1f1f1').grid(column=1,row=59)
Label(frame.scrollable_frame, text='                           ',font=('Arial',15),bg='#f1f1f1').grid(column=1,row=60)
Label(frame.scrollable_frame, text='                           ',font=('Arial',15),bg='#f1f1f1').grid(column=1,row=61)





#set focus to entry widget
city.focus()


#adding buttons
btn1 = Button(frame.scrollable_frame, text="Enter landlord details", bg="black", fg="white",command=lambda:get_landlords_details())
btn1.grid(column=2, row=14)



btn4 = Button(frame.scrollable_frame, text="Generate Agreement", 
              bg="orange", fg="white"
              ,command=lambda:createDoc
              (city.get(),state.get(),date.get(),
               landlord_details_final_global,
               tenant_details_final_global,prop.get(),
               var2,rent_am.get(),duration.get(),var4,exp_date.get()))
btn4.grid(column=3, row=100)

frame.grid(row=0, columnspan=500,rowspan=500,)
root.mainloop()






def createDoc(city,state,date,landlord_details_final,tenant_details_final,
              prop_address,rent_period,rent_amount,
              contract_duration,purpose_final,exp_date):
    doc=Document('editted1.docx')
    doc1=Document()
    fullText=[]    
    for para in doc.paragraphs:
        fullText.append(para.text)
    complete=('\n\n'.join(fullText))
    # editing place
    place=city+', '+state
    fo=complete.find('[PLACE]')
    complete1=complete[:fo]+place+complete[fo+len('[PLACE]'):]
    
    #editing date
    day=int(date[0]+date[1])
    year=int(date[6:])
    
    p = inflect.engine()
    lis_day=[]
    for i in range(1,32):
        lis_day.append(p.number_to_words(p.ordinal(i)).title())
        
        
    month_lst = ['January', 'February', 'March', 'April', 'May', 'June', 'July',
                  'August', 'September', 'October', 'November', 'December']
    
    month=int(date[3]+date[4])
    
    date_complete=str(date)+' ('+str(lis_day[day-1])+' day '+ 'of' + ' '+ str(month_lst[month-1])+', ' + p.number_to_words(year).title()+')'
    
    fo=complete1.find('[DatedayofMonth,Year]')
    complete2=complete1[:fo]+date_complete+complete1[fo+len('[DatedayofMonth,Year]'):]
    #adding name and details of landlord
    fo=complete2.find('[NameoftheLandlord], [bearingAadhar/CINnumber] [ResidentofPermanentAddressoftheLandlord] ')
    complete3=complete2[:fo]+'\n'+landlord_details_final+'\n'+complete2[fo+len('[NameoftheLandlord], [bearingAadhar/CINnumber] [ResidentofPermanentAddressoftheLandlord] '):]
    #jointlyandseverally
    if (int(numb_landlords))>1:
        fo=complete3.find('[jointlyandseverally]')
        complete4=complete3[:fo]+'jointly and severally'+complete3[fo+len('[jointlyandseverally]'):]
    else:
        fo=complete3.find('[jointlyandseverally]')
        complete4=complete3[:fo]+complete3[fo+len('[jointlyandseverally]'):]
    #ading tenants
    fo=complete4.find('[NameoftheTenant], having permanent address at [CompletepermanentAddressoftheTenant] and [bearingAadhar/CINnumber] having [PassportNumberissuedfromNameoftheCountryonDateofissuanceofPassport]')
    complete5=complete4[:fo]+'\n'+tenant_details_final+'\n'+complete4[fo+len('[NameoftheTenant], having permanent address at [CompletepermanentAddressoftheTenant] and [bearingAadhar/CINnumber] having [PassportNumberissuedfromNameoftheCountryonDateofissuanceofPassport]'):]
    #jointly and severally for tenant
    if (int(numb_tenants))>1:
        fo=complete5.find('[jointlyandseverally]')
        complete6=complete5[:fo]+'jointly and severally'+complete5[fo+len('[jointlyandseverally]'):]
    else:
        fo=complete5.find('[jointlyandseverally]')
        complete6=complete5[:fo]+complete5[fo+len('[jointlyandseverally]'):]
    #property address
    fo=complete6.find('[CompleteAddressoftheProperty]')
    complete7=complete6[:fo]+prop_address+complete6[fo+len('[CompleteAddressoftheProperty]'):]
    p= inflect.engine()
    rent_in_words = p.number_to_words(rent_amount).title()
    rent_final='INR '+str(rent_amount)+' (Indian Rupees '+ rent_in_words+') '+rent_period

    fo=complete7.find('[rentamountinINRpermonth/yearly/quarterly]')
    complete8=complete7[:fo]+rent_final+complete7[fo+len('[rentamountinINRpermonth/yearly/quarterly]'):]
    
    #contract duration
    p= inflect.engine()
    contract_duration_in_words=p.number_to_words(contract_duration).title()
    contract_duration_final=str(contract_duration)+' Months'+ ' (' +contract_duration_in_words+' Months)'

    fo=complete8.find('[inmonths]')
    complete9=complete8[:fo]+contract_duration_final+complete8[fo+len('[inmonths]'):]
    #enter purpose
    fo=complete9.find('[residential/commercialetc.]')
    complete10=complete9[:fo]+purpose_final+complete9[fo+len('[residential/commercialetc.]'):]
    
    #expiry date
    exp_day=int(exp_date[0]+exp_date[1])
    exp_year=int(exp_date[6:])

    exp_month=int(exp_date[3]+exp_date[4])

    exp_date_complete=str(exp_date)+' ('+str(lis_day[exp_day-1])+' day '+ 'of' + ' '+ str(month_lst[exp_month-1])+', ' + p.number_to_words(exp_year).title()+')'

    fo=complete10.find('[Expiry Date of Agreement]')
    complete11=complete10[:fo]+exp_date_complete+complete10[fo+len('[Expiry Date of Agreement]'):]


    doc1.add_paragraph(complete11)
    doc1.save('editted3.docx')
    